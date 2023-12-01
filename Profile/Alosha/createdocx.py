import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import sys

def createdocx(docxname, profile, path='/'):
    # Create a new Word document
    doc = Document()

    # Define a style for the title
    title_format = doc.styles.add_style('TitleFormat', 1)
    title_format.font.size = Pt(20)
    title_format.font.bold = True

    bodytitle_format = doc.styles.add_style('BodyTitleFormat', 1)
    bodytitle_format.font.size = Pt(14)
    bodytitle_format.font.bold = True

    # Define a style for the body
    body_format = doc.styles.add_style('BodyFormat', 1)
    body_format.font.size = Pt(12)


    # Add the person's name as a title
    doc.add_paragraph(profile['first_name']+' '+profile['last_name'], style='TitleFormat')

    # Add the person's contact information
    doc.add_paragraph(f"Country: {profile['country']}")
    doc.add_paragraph(f"Address: {profile['street']}, {profile['city']}, {profile['location']}, {profile['zipcode']}")
    doc.add_paragraph(f"Phone: {profile['phone']}")

    # Add a line break
    doc.add_paragraph('Overview', style='TitleFormat')
    doc.add_paragraph(profile['overview'], style='BodyFormat')

    # Add a line break
    doc.add_paragraph()

    # Add the person's work experience
    doc.add_paragraph('Work Experience', style='TitleFormat')
    for job in profile['workXP']:
        doc.add_paragraph(f"{job['role']}", style='BodyTitleFormat')
        doc.add_paragraph(f"{job['company']}, {job['start']} - {job['end']}", style='BodyFormat')
        for description in job['description']:
            doc.add_paragraph('• '+description, style='BodyFormat')
        doc.add_paragraph()

    # Add a line break
    doc.add_paragraph()

    # Add the person's education
    doc.add_paragraph('Education', style='TitleFormat')
    for education in profile['education']:
        doc.add_paragraph(f"{education['start']} - {education['end']}", style='BodyFormat')
        doc.add_paragraph(f"{education['degree']} in {education['field']}", style='BodyTitleFormat')
        doc.add_paragraph(f"{education['university']}, {education['country']}", style='BodyFormat')

    # Add a line break
    doc.add_paragraph()


    # # Add the person's skills
    # doc.add_paragraph('Skills', style='TitleFormat')
    # doc.add_paragraph(', '.join(profile['skills']), style='BodyFormat')

    # # Add a line break
    # doc.add_paragraph()

    # Add the person's language
    doc.add_paragraph('Language', style='TitleFormat')
    if len(profile['languages'])==0:
        doc.add_paragraph('English - Fluent', style='BodyFormat')
    else:
        doc.add_paragraph(profile['languages'][0]['language']+'-'+profile['languages'][0]['level'], style='BodyFormat')

    # Save the Word document
    # docxpath = filename.capitalize() + '.docx'
    docxpath = docxname + '.docx'
    # docxpath = path + docxname + '.docx'
    doc.save(docxpath)
    print("INFO: Docx resume file is created on " + docxpath)
    return docxpath

if __name__ == '__main__':
    # Load the JSON profile data
    # filename = sys.argv[1]
    files = os.listdir()# Filter files with .json extension
    json_files = [file for file in files if file.endswith('.json')]

    for i in json_files:
        with open(i, 'r', encoding='utf-8') as f:
            profile = json.load(f)
        createdocx(i.replace('.json',''), profile)
